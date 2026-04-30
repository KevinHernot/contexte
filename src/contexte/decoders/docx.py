"""DOCX decoder."""

from __future__ import annotations

from pathlib import Path
from typing import Any, ClassVar

from contexte.decoders.base import BaseDecoder, DecodeContext, read_text_with_fallback
from contexte.ir.models import ContextDocument, ExtractionError
from contexte.normalizers.text import normalize_text, split_paragraphs


class DocxDecoder(BaseDecoder):
    id = "docx"
    supported_extensions: ClassVar[set[str]] = {".docx"}
    supported_media_types: ClassVar[set[str]] = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }

    def decode(self, path: Path, context: DecodeContext) -> ContextDocument:
        blocks = []
        errors: list[ExtractionError] = []
        index = 0
        title = path.stem
        try:
            from docx import Document

            document = Document(str(path))
            for paragraph in document.paragraphs:
                text = normalize_text(paragraph.text)
                if not text:
                    continue
                style = (paragraph.style.name if paragraph.style is not None else "").lower()
                if style.startswith("heading"):
                    block_type = "heading"
                    level = _heading_level(style)
                    if title == path.stem and level == 1:
                        title = text
                else:
                    block_type = "paragraph"
                    level = None
                blocks.append(self.block(context, index, block_type, text, level=level))
                index += 1
            for table in document.tables:
                markdown = _table_to_markdown(table)
                text = normalize_text(markdown)
                if text:
                    blocks.append(
                        self.block(
                            context, index, "table", text, markdown=markdown, confidence=0.85
                        )
                    )
                    index += 1
        except Exception as exc:
            errors.append(
                ExtractionError(
                    code="docx_decode_failed",
                    message=f"python-docx could not decode the file: {exc}",
                    severity="warning",
                )
            )
            raw, encoding = read_text_with_fallback(path)
            for fallback_paragraph in split_paragraphs(raw):
                blocks.append(
                    self.block(
                        context,
                        index,
                        "paragraph",
                        normalize_text(fallback_paragraph),
                        confidence=0.25,
                    )
                )
                index += 1
            if blocks:
                errors.append(
                    ExtractionError(
                        code="docx_text_fallback",
                        message=f"Used low-confidence text fallback with {encoding} decoding.",
                        severity="warning",
                    )
                )

        if not blocks:
            errors.append(
                ExtractionError(
                    code="docx_empty",
                    message="No text or table content was extracted from the DOCX file.",
                    severity="warning",
                )
            )
        return self.make_document(path, context, blocks, title=title, errors=errors)


def _heading_level(style_name: str) -> int | None:
    parts = style_name.split()
    for part in reversed(parts):
        if part.isdigit():
            return max(1, min(6, int(part)))
    return 1


def _table_to_markdown(table: Any) -> str:
    rows = [
        [normalize_text(cell.text, preserve_line_breaks=False) for cell in row.cells]
        for row in table.rows
    ]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)
